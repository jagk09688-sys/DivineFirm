from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Config: update these values if you have the real ABN/address
COMPANY_NAME = "DF Divine Firm"
ABN = "ABN: 90701540338"
ADDRESS = "24 Borthwick Parade, Clyde North"
EMAIL = "divinefirm30@gmail.com"
PHONE = ""
WEBSITE = "www.divinefirm.online"

out_path = "assets/divinefirm-letterhead.docx"

doc = Document()
section = doc.sections[0]

# Header
header = section.header
h_p = header.paragraphs[0]
h_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = h_p.add_run(COMPANY_NAME + "\n")
run.bold = True
run.font.size = Pt(20)

h_p2 = header.add_paragraph(f"{ABN} | {ADDRESS}")
h_p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
h_p2.runs[0].font.size = Pt(9)

h_p3 = header.add_paragraph(f"Email: {EMAIL} {(' | Phone: ' + PHONE) if PHONE else ''}")
h_p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
h_p3.runs[0].font.size = Pt(9)

# Add a thin separator (a paragraph with bottom border)
sep = doc.add_paragraph()
# Workaround: add a run of underscores as a divider
sep.add_run('_' * 100)

# Body template
doc.add_paragraph('\n')
doc.add_paragraph('Date: ______________________')
doc.add_paragraph('\n')
doc.add_paragraph('To:')
doc.add_paragraph('\n')
doc.add_paragraph('Subject:')
doc.add_paragraph('\n')
doc.add_paragraph('Dear Sir/Madam,')
doc.add_paragraph('\n')
doc.add_paragraph('')

doc.add_paragraph('\n\nYours sincerely,')
doc.add_paragraph('\n')
doc.add_paragraph(COMPANY_NAME)

# Footer
footer = section.footer
f_p = footer.paragraphs[0]
f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
f_run = f_p.add_run(f"{COMPANY_NAME} • {ABN} • {EMAIL} • {WEBSITE}")
f_run.font.size = Pt(9)

# Save
doc.save(out_path)
print(f"Generated {out_path}")
